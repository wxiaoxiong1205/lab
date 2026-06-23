"""
评估报告DOCX生成器

用于将评估任务的基本信息和评估结果导出为DOCX格式文件，包含文本内容和图表。
"""
from datetime import datetime
from typing import List, Dict, Any, Optional
from io import BytesIO
import math

import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.evaluation_task import (
    EvaluationTaskDetailResponse,
    EvaluationReportResponse,
    ModelReportData,
    AggregativeMetric,
    ModelMetricSummary,
    EvaluationMethod,
    EvaluationType
)
from app.core.logging import logger
from app.utils.pdf_utils import get_system_chinese_font
from app.utils.docx_font_utils import (
    apply_east_asia_font_to_all_runs,
    setup_document_east_asia_font,
    set_run_east_asia_font,
)

# 配置matplotlib中文字体：优先使用系统实际存在的中文字体文件，避免 Linux 等环境仅剩 DejaVu Sans 导致中文乱码
_plt_font_sans_serif = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'DejaVu Sans']
_chinese_font_path = get_system_chinese_font()
if _chinese_font_path:
    try:
        if hasattr(fm.fontManager, 'addfont'):
            fm.fontManager.addfont(_chinese_font_path)
        _prop = fm.FontProperties(fname=_chinese_font_path)
        _chinese_font_name = _prop.get_name()
        if _chinese_font_name:
            _plt_font_sans_serif = [_chinese_font_name] + _plt_font_sans_serif
            logger.debug(f"matplotlib 中文字体已设置为: {_chinese_font_name} ({_chinese_font_path})")
    except Exception as e:
        logger.warning(f"matplotlib 中文字体注册失败，图表中文可能乱码: {e}")
plt.rcParams['font.sans-serif'] = _plt_font_sans_serif
plt.rcParams['axes.unicode_minus'] = False


class EvaluationReportDocxGenerator:
    """评估报告DOCX生成器"""
    
    def __init__(
        self,
        task_data: EvaluationTaskDetailResponse,
        report_data: EvaluationReportResponse
    ):
        """初始化生成器
        
        Args:
            task_data: 评估任务基本信息
            report_data: 评估报告数据（包含各模型的指标汇总）
        """
        self.task_data = task_data
        self.report_data = report_data
        self.doc = Document()
        
        # 设置文档默认字体
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """设置文档样式：正文与标题统一为黑色黑体（图表为 matplotlib 生成，不受此影响）。"""
        setup_document_east_asia_font(self.doc)
        self._set_heading_styles_to_black()

    def _set_heading_styles_to_black(self):
        """将标题样式改为黑色，避免使用 Word 默认主题色（蓝色）。"""
        heading_style_names = ['Title'] + [f'Heading {level}' for level in range(1, 10)]
        for style_name in heading_style_names:
            try:
                style = self.doc.styles[style_name]
            except KeyError:
                continue

            font = getattr(style, 'font', None)
            if font is None:
                continue

            font.name = '黑体'
            font.color.rgb = None

            if style_name == 'Title':
                font.bold = True
                font.size = Pt(18)
            else:
                level = int(style_name.split()[-1])
                if level == 1:
                    font.bold = True
                    font.size = Pt(16)
                elif level == 2:
                    font.bold = True
                    font.size = Pt(14)
                elif level == 3:
                    font.bold = True
                    font.size = Pt(12)

        for paragraph in self.doc.paragraphs:
            style_name = getattr(paragraph.style, 'name', '')
            if style_name in heading_style_names:
                for run in paragraph.runs:
                    set_run_east_asia_font(run)
                    run.font.color.rgb = None
    
    def generate(self) -> bytes:
        """生成DOCX文件并返回字节流
        
        Returns:
            DOCX文件的字节流
        """
        try:
            # 1. 添加标题
            self._add_title()
            
            # 2. 添加基本信息章节
            self._add_basic_info()

            # 2. 添加评估结果章节（二、报告章节（平均））
            self._add_model_results()

            # 3. 添加评分对比柱状图章节（三、评分对比柱状图）
            self._add_bar_charts()

            # 4. 正文字体统一为黑体后保存
            apply_east_asia_font_to_all_runs(self.doc)
            buffer = BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成DOCX文件失败: {str(e)}", exc_info=True)
            raise
    
    def _add_title(self):
        """添加报告标题"""
        # 标题
        title = self.doc.add_heading('评估报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = title._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn('w:pBdr'))
        if p_bdr is None:
            p_bdr = OxmlElement('w:pBdr')
            p_pr.append(p_bdr)
        bottom = p_bdr.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            p_bdr.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        
        # 任务名称
        task_name_para = self.doc.add_paragraph()
        task_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = task_name_para.add_run(self.task_data.name)
        run.font.size = Pt(16)
        run.bold = True
        
        # 生成时间
        time_para = self.doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_str = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        time_para.add_run(f"生成时间：{time_str}")
        
        # 添加分页
        self.doc.add_page_break()
    
    def _add_basic_info(self):
        """添加基本信息章节（10个字段，表格格式）"""
        self.doc.add_heading('一、任务基本信息', 1)
        
        # 待评估模型/服务：从 dataset_model_relations 获取
        evaluated_models = []
        inference_datasets = []
        if self.task_data.dataset_model_relations:
            for rel in self.task_data.dataset_model_relations:
                if rel.evaluated_model_name:
                    evaluated_models.append(rel.evaluated_model_name)
                if rel.inference_result_dataset_name:
                    inference_datasets.append(rel.inference_result_dataset_name)
        evaluated_models_str = '，'.join(evaluated_models) if evaluated_models else '无'
        inference_datasets_str = '，'.join(inference_datasets) if inference_datasets else '无'
        
        # 评估方法：all 时列出两种，referee 时只裁判员，basic_metric 时只基础指标
        if self.task_data.evaluation_method == EvaluationMethod.ALL:
            evaluation_method_str = '裁判员评估，基础指标评估'
        elif self.task_data.evaluation_method == EvaluationMethod.REFEREE:
            evaluation_method_str = '裁判员评估'
        elif self.task_data.evaluation_method == EvaluationMethod.BASIC_METRIC:
            evaluation_method_str = '基础指标评估'
        else:
            evaluation_method_str = self._get_evaluation_method_name(self.task_data.evaluation_method)
        
        # 裁判员模型/服务：referee_model_name + "/" + (离线/在线)，无裁判员时显示"无"
        referee_str = '无'
        if self.task_data.evaluation_method in [EvaluationMethod.REFEREE, EvaluationMethod.ALL]:
            if self.task_data.referee_model_name:
                referee_type_name = '离线' if self.task_data.referee_type == 'model' else '在线'
                referee_str = f"{self.task_data.referee_model_name}/{referee_type_name}"
        
        # 评估类别：dataset_type 映射
        dataset_type_str = self._get_dataset_type_name(self.task_data.dataset_type)
        
        # 创建时间
        created_at_str = self.task_data.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.task_data.created_at else '无'
        
        # 描述
        description_str = self.task_data.description or '无'
        
        # 基本信息字段（10个，无序号）
        basic_info_fields = [
            ('任务名称', self.task_data.name),
            ('待评估模型/服务', evaluated_models_str),
            ('推理结果集', inference_datasets_str),
            ('评估类型', self._get_evaluation_type_name(self.task_data.evaluation_type)),
            ('评估类别', dataset_type_str),
            ('评估方法', evaluation_method_str),
            ('裁判员模型/服务', referee_str),
            ('创建人', self.task_data.created_by or '无'),
            ('创建时间', created_at_str),
            ('描述', description_str),
        ]
        
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for field_name, field_value in basic_info_fields:
            row = table.add_row()
            row.cells[0].text = field_name
            row.cells[1].text = str(field_value)
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        self.doc.add_paragraph()
    
    def _add_model_results(self):
        """添加评估结果章节：二、报告章节（平均）"""
        self.doc.add_heading('二、报告章节（平均）', 1)
        
        if not self.report_data.model_reports:
            self.doc.add_paragraph("暂无评估结果数据。")
            return
        
        referee_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.REFEREE]
        basic_metric_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.BASIC_METRIC]
        manual_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.MANUAL]

        # 1. 裁判员评估
        if referee_reports:
            self.doc.add_heading('1. 裁判员评估', 2)
            # 1.1 评分维度雷达图
            self.doc.add_heading('1.1. 评分维度雷达图', 3)
            metric_names = self._get_referee_metric_names()
            if metric_names:
                try:
                    radar_image = self._generate_radar_chart(referee_reports, metric_names)
                    if radar_image:
                        para = self.doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(BytesIO(radar_image), width=Inches(6))
                except Exception as e:
                    logger.warning(f"生成裁判员雷达图失败: {str(e)}")
                    self.doc.add_paragraph(f"[雷达图生成失败：{str(e)}]")
            # 1.2 评估数据明细
            self.doc.add_heading('1.2. 评估数据明细', 3)
            self._add_evaluation_detail_table(referee_reports)
            self.doc.add_paragraph()
        
        # 2. 基础指标评估
        if basic_metric_reports:
            self.doc.add_heading('2. 基础指标评估', 2)
            # 2.1 评分维度雷达图
            self.doc.add_heading('2.1. 评分维度雷达图', 3)
            metric_names = self._get_basic_metric_names()
            if metric_names:
                try:
                    radar_image = self._generate_radar_chart(basic_metric_reports, metric_names)
                    if radar_image:
                        para = self.doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(BytesIO(radar_image), width=Inches(6))
                except Exception as e:
                    logger.warning(f"生成基础指标雷达图失败: {str(e)}")
                    self.doc.add_paragraph(f"[雷达图生成失败：{str(e)}]")
            # 2.2 评估数据明细
            self.doc.add_heading('2.2. 评估数据明细', 3)
            self._add_evaluation_detail_table(basic_metric_reports)
            self.doc.add_paragraph()

        # 3. 人工指标评估
        if manual_reports  and not referee_reports and not basic_metric_reports:
            self.doc.add_heading('1. 人工评估', 2)
            # 2.1 评分维度雷达图
            self.doc.add_heading('1.1. 评分维度雷达图', 3)
            metric_names = self._get_manual_metric_names()
            if metric_names:
                try:
                    radar_image = self._generate_radar_chart(manual_reports, metric_names)
                    if radar_image:
                        para = self.doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(BytesIO(radar_image), width=Inches(6))
                except Exception as e:
                    logger.warning(f"生成基础指标雷达图失败: {str(e)}")
                    self.doc.add_paragraph(f"[雷达图生成失败：{str(e)}]")
            # 2.2 评估数据明细
            self.doc.add_heading('2.2. 评估数据明细', 3)
            self._add_evaluation_detail_table(manual_reports)
            self.doc.add_paragraph()

    def _get_referee_metric_names(self) -> List[str]:
        """从评估配置获取裁判员指标名称"""
        if self.task_data.evaluation_prompt_config and self.task_data.evaluation_prompt_config.metrics:
            return [m.name for m in self.task_data.evaluation_prompt_config.metrics]
        # 从报告数据中获取（保持接口返回顺序）
        for r in self.report_data.model_reports:
            if r.evaluation_method == EvaluationMethod.REFEREE and r.aggregative_metrics:
                for agg in r.aggregative_metrics:
                    if agg.metric_summary:
                        return list(agg.metric_summary.keys())
        return []
    
    def _get_basic_metric_names(self) -> List[str]:
        """从评估配置获取基础指标名称"""
        if self.task_data.basic_metric_config and self.task_data.basic_metric_config.metrics:
            return self.task_data.basic_metric_config.metrics
        for r in self.report_data.model_reports:
            if r.evaluation_method == EvaluationMethod.BASIC_METRIC and r.aggregative_metrics:
                for agg in r.aggregative_metrics:
                    if agg.metric_summary:
                        return list(agg.metric_summary.keys())
        return []

    def _get_manual_metric_names(self) -> List[str]:
        """从评估配置获取裁判员指标名称"""
        if self.task_data.evaluation_prompt_config and self.task_data.evaluation_prompt_config.metrics:
            return [m.name for m in self.task_data.evaluation_prompt_config.metrics]
        # 从报告数据中获取（保持接口返回顺序）
        for r in self.report_data.model_reports:
            if r.evaluation_method == EvaluationMethod.MANUAL and r.aggregative_metrics:
                for agg in r.aggregative_metrics:
                    if agg.metric_summary:
                        return list(agg.metric_summary.keys())
        return []

    def _add_evaluation_detail_table(self, model_reports: List[ModelReportData]):
        """添加评估数据明细表（评估指标 | 模型1 | 模型2 | 模型3，百分比得分，表头加粗）"""
        if not model_reports or not model_reports[0].aggregative_metrics:
            self.doc.add_paragraph("暂无评估数据。")
            return
        agg = model_reports[0].aggregative_metrics[0]
        metric_names = list(agg.metric_summary.keys())  # 按接口返回顺序，与 benchmark_compare 一致
        model_names = [r.model_name for r in model_reports]
        
        table = self.doc.add_table(rows=1, cols=len(model_names) + 1)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '评估指标'
        for i, mn in enumerate(model_names):
            header_cells[i + 1].text = mn
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for metric_name in metric_names:
            row = table.add_row()
            row.cells[0].text = metric_name
            for i, model_report in enumerate(model_reports):
                metric_summary = None
                for a in model_report.aggregative_metrics:
                    if a.metric_summary and metric_name in a.metric_summary:
                        metric_summary = a.metric_summary[metric_name]
                        break
                val = f"{metric_summary.percentage_score:.2f}" if metric_summary and metric_summary.percentage_score is not None else "0.00"
                row.cells[i + 1].text = val
    
    def _add_bar_charts(self):
        """添加第三章节：三、评分对比柱状图"""
        self.doc.add_heading('三、评分对比柱状图', 1)
        
        if not self.report_data.model_reports:
            self.doc.add_paragraph("暂无评估结果数据。")
            return
        
        referee_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.REFEREE]
        basic_metric_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.BASIC_METRIC]
        manual_reports = [r for r in self.report_data.model_reports if r.evaluation_method == EvaluationMethod.MANUAL]

        if referee_reports:
            self.doc.add_heading('1. 裁判员评估', 2)
            try:
                bar_image = self._generate_comparison_bar_chart(referee_reports)
                if bar_image:
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(BytesIO(bar_image), width=Inches(6))
            except Exception as e:
                logger.warning(f"生成裁判员柱状图失败: {str(e)}")
                self.doc.add_paragraph(f"[柱状图生成失败：{str(e)}]")
            self.doc.add_paragraph()
        
        if basic_metric_reports:
            self.doc.add_heading('2. 基础指标评估', 2)
            try:
                bar_image = self._generate_comparison_bar_chart(basic_metric_reports)
                if bar_image:
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(BytesIO(bar_image), width=Inches(6))
            except Exception as e:
                logger.warning(f"生成基础指标柱状图失败: {str(e)}")
                self.doc.add_paragraph(f"[柱状图生成失败：{str(e)}]")

        if manual_reports  and not referee_reports and not basic_metric_reports:
            self.doc.add_heading('1. 手工评估', 2)
            try:
                bar_image = self._generate_comparison_bar_chart(manual_reports)
                if bar_image:
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(BytesIO(bar_image), width=Inches(6))
            except Exception as e:
                logger.warning(f"生成基础指标柱状图失败: {str(e)}")
                self.doc.add_paragraph(f"[柱状图生成失败：{str(e)}]")

    def _generate_radar_chart(self, model_reports: List[ModelReportData], metric_names: List[str]) -> Optional[bytes]:
        """生成雷达图（多边形样式）：每个顶点为指标，不同颜色的线代表不同模型，使用百分比得分。
        一个指标为一条线，三个指标为三角形，四个为四边形，依次类推。"""
        if not metric_names or not model_reports:
            return None
        try:
            num_metrics = len(metric_names)
            # 顶点角度：从顶部(pi/2)开始，顺时针。一个指标时为一条线
            if num_metrics == 1:
                vertex_angles = [math.pi / 2]  # 顶点在正上方
            else:
                vertex_angles = [2 * math.pi * i / num_metrics - math.pi / 2 for i in range(num_metrics)]
            
            fig, ax = plt.subplots(figsize=(8, 8))
            colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', '#13c2c2', '#eb2f96']
            max_r = 1.0
            
            # 绘制多边形网格（同心多边形）
            for level in [0.25, 0.5, 0.75, 1.0]:
                r = max_r * level
                if num_metrics == 1:
                    xs, ys = [0, r * math.cos(vertex_angles[0])], [0, r * math.sin(vertex_angles[0])]
                else:
                    poly_angles = vertex_angles + [vertex_angles[0]]
                    xs = [r * math.cos(a) for a in poly_angles]
                    ys = [r * math.sin(a) for a in poly_angles]
                ax.plot(xs, ys, 'k-', linewidth=0.5, alpha=0.5)
            
            # 绘制从中心到各顶点的射线
            for a in vertex_angles:
                ax.plot([0, max_r * math.cos(a)], [0, max_r * math.sin(a)], 'k-', linewidth=0.5, alpha=0.5)
            
            # 添加刻度标签（25, 50, 75, 100），在每条射线上标注
            offset = 0.04
            for tick_angle in vertex_angles:
                for level, pct in [(0.25, 25), (0.5, 50), (0.75, 75), (1.0, 100)]:
                    r = max_r * level
                    base_x = r * math.cos(tick_angle)
                    base_y = r * math.sin(tick_angle)
                    nx, ny = math.cos(tick_angle + math.pi / 2), math.sin(tick_angle + math.pi / 2)
                    x = base_x + offset * nx
                    y = base_y + offset * ny
                    ax.text(x, y, f'{pct:.2f}', ha='center', va='center', fontsize=9)
            
            # 绘制各模型数据（多边形或线）
            for idx, model_report in enumerate(model_reports):
                values = []
                for mn in metric_names:
                    metric_summary = None
                    for agg in model_report.aggregative_metrics or []:
                        if agg.metric_summary and mn in agg.metric_summary:
                            metric_summary = agg.metric_summary[mn]
                            break
                    val = (metric_summary.percentage_score or 0) / 100.0 if metric_summary else 0
                    values.append(val)
                radii = [max_r * v for v in values]
                if num_metrics == 1:
                    xs = [0, radii[0] * math.cos(vertex_angles[0])]
                    ys = [0, radii[0] * math.sin(vertex_angles[0])]
                    ax.plot(xs, ys, 'o-', linewidth=2, label=model_report.model_name, color=colors[idx % len(colors)])
                else:
                    poly_angles = vertex_angles + [vertex_angles[0]]
                    poly_radii = radii + [radii[0]]
                    xs = [poly_radii[i] * math.cos(poly_angles[i]) for i in range(len(poly_angles))]
                    ys = [poly_radii[i] * math.sin(poly_angles[i]) for i in range(len(poly_angles))]
                    ax.plot(xs, ys, 'o-', linewidth=2, label=model_report.model_name, color=colors[idx % len(colors)])
                    ax.fill(xs, ys, alpha=0.15, color=colors[idx % len(colors)])
            
            # 顶点标签
            label_r = max_r * 1.08
            for i in range(num_metrics):
                a = vertex_angles[i]
                ax.text(label_r * math.cos(a), label_r * math.sin(a), metric_names[i],
                        ha='center', va='center', fontsize=10)
            
            ax.set_xlim(-1.25 * max_r, 1.25 * max_r)
            ax.set_ylim(-1.25 * max_r, 1.25 * max_r)
            ax.set_aspect('equal')
            ax.axis('off')
            ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.0))
            plt.tight_layout()
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成雷达图失败: {str(e)}", exc_info=True)
            plt.close()
            return None
    
    def _generate_comparison_bar_chart(self, model_reports: List[ModelReportData]) -> Optional[bytes]:
        """生成对比柱状图：横坐标为指标，纵坐标为原始得分score，不同颜色柱体代表不同模型"""
        if not model_reports or not model_reports[0].aggregative_metrics:
            return None
        try:
            agg = model_reports[0].aggregative_metrics[0]
            metric_names = list(agg.metric_summary.keys())  # 按接口返回顺序
            model_names = [r.model_name for r in model_reports]
            
            x = range(len(metric_names))
            width = 0.8 / len(model_names)
            colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', '#13c2c2', '#eb2f96']
            
            fig, ax = plt.subplots(figsize=(12, 6))
            max_val = 0
            for i, model_report in enumerate(model_reports):
                scores = []
                for mn in metric_names:
                    metric_summary = None
                    for a in model_report.aggregative_metrics or []:
                        if a.metric_summary and mn in a.metric_summary:
                            metric_summary = a.metric_summary[mn]
                            break
                    s = (metric_summary.percentage_score or 0) if metric_summary else 0
                    scores.append(s)
                    max_val = max(max_val, s)
                offset = (i - len(model_names) / 2 + 0.5) * width
                bars = ax.bar([xi + offset for xi in x], scores, width, label=model_report.model_name, color=colors[i % len(colors)])
                for bar in bars:
                    h = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., h, f'{h:.2f}%', ha='center', va='bottom', fontsize=8)
            
            ax.set_ylabel('百分比得分 (%)', fontsize=12)
            ax.set_xticks(x)
            ax.set_xticklabels(metric_names, rotation=45, ha='right')
            ax.set_ylim(0, max(max_val * 1.2, 100) if max_val > 0 else 100)
            ax.legend()
            plt.tight_layout()
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成柱状图失败: {str(e)}", exc_info=True)
            plt.close()
            return None
    
    def _get_evaluation_type_name(self, evaluation_type: EvaluationType) -> str:
        """获取评估类型中文名称"""
        type_map = {
            EvaluationType.SINGLE: "单个评估",
            EvaluationType.COMPARISON: "对比评估"
        }
        return type_map.get(evaluation_type, str(evaluation_type))
    
    def _get_evaluation_method_name(self, evaluation_method: EvaluationMethod) -> str:
        """获取评估方法中文名称"""
        method_map = {
            EvaluationMethod.REFEREE: "裁判员评估",
            EvaluationMethod.BASIC_METRIC: "基础指标评估",
            EvaluationMethod.ALL: "同时评估",
            EvaluationMethod.MANUAL: "人工评估"
        }
        return method_map.get(evaluation_method, str(evaluation_method))
    
    def _get_data_source_name(self, data_source) -> str:
        """获取数据来源中文名称"""
        source_map = {
            "existing": "已有推理结果集",
            "new": "新建推理结果集"
        }
        if hasattr(data_source, 'value'):
            return source_map.get(data_source.value, str(data_source))
        return source_map.get(str(data_source), str(data_source))

    def _get_dataset_type_name(self, dataset_type: Optional[str]) -> str:
        """获取评估类别中文名称"""
        if not dataset_type:
            return '未设置'
        type_map = {
            "text-generation": "文本生成",
            "image-generation": "图像生成",
            "image-understanding": "图像理解",
            "multimodal": "多模态",
            "business": "文本生成",
        }
        return type_map.get(dataset_type, dataset_type)
