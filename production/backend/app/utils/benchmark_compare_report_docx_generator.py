"""
对比评估报告 DOCX 生成器

用于将基准对比评估接口返回的数据导出为 DOCX，包含：标题、评分维度雷达图、评估指标明细表、评分对比柱状图。
数据来源：POST /api/v1/benchmark/project/{project_id}/tasks/compare 的响应。
结构参照 benchmark_report_docx_generator，图表需明确坐标轴与刻度，并复用中文字体配置避免乱码。
"""
from datetime import datetime
from typing import List, Optional
from io import BytesIO
import math

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.benchmark_task import (
    BenchmarkTaskCompareResponse,
    BenchmarkModelReportData,
)
from app.core.logging import logger
from app.utils.pdf_utils import get_system_chinese_font
from app.utils.docx_font_utils import (
    apply_east_asia_font_to_all_runs,
    setup_document_east_asia_font,
    set_run_east_asia_font,
)

# 配置 matplotlib 中文字体（与基准报告一致，避免图表中文乱码）
_plt_font_sans_serif = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'DejaVu Sans']
_chinese_font_path = get_system_chinese_font()
if _chinese_font_path:
    try:
        if hasattr(fm.fontManager, 'addfont'):
            fm.fontManager.addfont(_chinese_font_path)
        _prop = fm.FontProperties(fname=_chinese_font_path)
        _chinese_font_name = _prop.get_name()
        if _chinese_font_name:
            _plt_font_sans_serif = [_chinese_font_name] + _plt_font_sans_serif
            logger.debug(f"matplotlib 中文字体已设置为: {_chinese_font_name} ({_chinese_font_path})")
    except Exception as e:
        logger.warning(f"matplotlib 中文字体注册失败，图表中文可能乱码: {e}")
plt.rcParams['font.sans-serif'] = _plt_font_sans_serif
plt.rcParams['axes.unicode_minus'] = False


class BenchmarkCompareReportDocxGenerator:
    """对比评估报告 DOCX 生成器。仅包含：标题、1. 对比评估（雷达图 + 评估指标明细）、三、评分对比柱状图。"""

    def __init__(self, compare_data: BenchmarkTaskCompareResponse):
        """
        Args:
            compare_data: 对比评估接口返回数据（model_reports 中每项含 radar_chart_data.data）。
        """
        self.compare_data = compare_data
        self.doc = Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        """设置文档样式：正文与标题统一为黑色黑体（图表为 matplotlib 生成，不受此影响）。"""
        setup_document_east_asia_font(self.doc)
        self._set_heading_styles_to_black()

    def _set_heading_styles_to_black(self):
        """将标题样式改为黑色，避免使用 Word 默认主题色（蓝色）。"""
        heading_style_names = ['Title'] + [f'Heading {level}' for level in range(1, 10)]
        for style_name in heading_style_names:
            try:
                style = self.doc.styles[style_name]
            except KeyError:
                continue

            font = getattr(style, 'font', None)
            if font is None:
                continue

            font.name = '黑体'
            font.color.rgb = None

            if style_name == 'Title':
                font.bold = True
                font.size = Pt(18)
            else:
                level = int(style_name.split()[-1])
                if level == 1:
                    font.bold = True
                    font.size = Pt(16)
                elif level == 2:
                    font.bold = True
                    font.size = Pt(14)
                elif level == 3:
                    font.bold = True
                    font.size = Pt(12)

        for paragraph in self.doc.paragraphs:
            style_name = getattr(paragraph.style, 'name', '')
            if style_name in heading_style_names:
                for run in paragraph.runs:
                    set_run_east_asia_font(run)
                    run.font.color.rgb = None

    def generate(self) -> bytes:
        """生成 DOCX 并返回字节流"""
        try:
            self._add_title()
            self._add_compare_section()
            self._add_bar_charts()
            apply_east_asia_font_to_all_runs(self.doc)
            buffer = BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成对比评估DOCX失败: {str(e)}", exc_info=True)
            raise

    def _add_title(self):
        """第一部分：标题（与其他报告保持一致）"""
        title = self.doc.add_heading('对比评估报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = title._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn('w:pBdr'))
        if p_bdr is None:
            p_bdr = OxmlElement('w:pBdr')
            p_pr.append(p_bdr)
        bottom = p_bdr.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            p_bdr.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')

        time_para = self.doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        self.doc.add_page_break()

    def _get_dataset_codes_and_labels(self) -> tuple:
        """从 model_reports 中取有序的 dataset_code 及对应展示名（dataset_name）。"""
        if not self.compare_data.model_reports:
            return [], []
        first = self.compare_data.model_reports[0]
        data = first.radar_chart_data.data if first.radar_chart_data else []
        codes = [item.dataset_code for item in data]
        labels = [item.dataset_name or item.dataset_code for item in data]
        return codes, labels

    def _model_report_to_dataset_scores(self, report: BenchmarkModelReportData) -> dict:
        """将单个 model_report 的 radar_chart_data.data 转为 dataset_code -> score。"""
        scores = {}
        if report.radar_chart_data and report.radar_chart_data.data:
            for item in report.radar_chart_data.data:
                if item.score is not None:
                    scores[item.dataset_code] = item.score
        return scores

    def _add_compare_section(self):
        """第二部分：1. 对比评估 -> 1.1 评分维度雷达图、1.2 评估指标明细"""
        if not self.compare_data.model_reports:
            self.doc.add_paragraph("暂无对比评估数据。")
            return
        dataset_codes, metric_labels = self._get_dataset_codes_and_labels()
        if not dataset_codes:
            self.doc.add_paragraph("暂无评估指标数据。")
            return

        self.doc.add_heading('1. 对比评估', 2)
        self.doc.add_heading('1.1. 评分维度雷达图', 3)
        try:
            radar_image = self._generate_radar_chart(dataset_codes, metric_labels)
            if radar_image:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(BytesIO(radar_image), width=Inches(6))
        except Exception as e:
            logger.warning(f"生成对比评估雷达图失败: {str(e)}")
            self.doc.add_paragraph(f"[雷达图生成失败：{str(e)}]")
        self.doc.add_heading('1.2. 评估指标明细', 3)
        self._add_detail_table(dataset_codes, metric_labels)
        self.doc.add_paragraph()

    def _add_detail_table(self, dataset_codes: List[str], metric_labels: List[str]):
        """评估指标明细表：表头为【评估指标、模型1、模型2…】"""
        model_names = [r.model_name for r in self.compare_data.model_reports]
        table = self.doc.add_table(rows=1, cols=len(model_names) + 1)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '评估指标'
        for i, name in enumerate(model_names):
            header_cells[i + 1].text = name
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for idx, code in enumerate(dataset_codes):
            row = table.add_row()
            row.cells[0].text = metric_labels[idx]
            for i, model_report in enumerate(self.compare_data.model_reports):
                scores = self._model_report_to_dataset_scores(model_report)
                score = scores.get(code)
                row.cells[i + 1].text = f"{score:.2f}" if score is not None else "-"
        self.doc.add_paragraph()

    def _generate_radar_chart(self, dataset_codes: List[str], metric_labels: List[str]) -> Optional[bytes]:
        """雷达图：顶点为评估指标，不同颜色的线代表不同模型，得分 0-100。"""
        if not dataset_codes or not self.compare_data.model_reports:
            return None
        try:
            num_metrics = len(dataset_codes)
            if num_metrics == 1:
                vertex_angles = [math.pi / 2]
            else:
                vertex_angles = [2 * math.pi * i / num_metrics - math.pi / 2 for i in range(num_metrics)]
            fig, ax = plt.subplots(figsize=(8, 8))
            colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', '#13c2c2', '#eb2f96']
            max_r = 1.0
            for level in [0.25, 0.5, 0.75, 1.0]:
                r = max_r * level
                if num_metrics == 1:
                    xs = [0, r * math.cos(vertex_angles[0])]
                    ys = [0, r * math.sin(vertex_angles[0])]
                else:
                    poly_angles = vertex_angles + [vertex_angles[0]]
                    xs = [r * math.cos(a) for a in poly_angles]
                    ys = [r * math.sin(a) for a in poly_angles]
                ax.plot(xs, ys, 'k-', linewidth=0.5, alpha=0.5)
            for a in vertex_angles:
                ax.plot([0, max_r * math.cos(a)], [0, max_r * math.sin(a)], 'k-', linewidth=0.5, alpha=0.5)
            offset = 0.04
            for tick_angle in vertex_angles:
                for level, pct in [(0.25, 25), (0.5, 50), (0.75, 75), (1.0, 100)]:
                    r = max_r * level
                    base_x = r * math.cos(tick_angle)
                    base_y = r * math.sin(tick_angle)
                    nx = math.cos(tick_angle + math.pi / 2)
                    ny = math.sin(tick_angle + math.pi / 2)
                    x = base_x + offset * nx
                    y = base_y + offset * ny
                    ax.text(x, y, f'{pct:.2f}', ha='center', va='center', fontsize=9)
            for idx, model_report in enumerate(self.compare_data.model_reports):
                scores = self._model_report_to_dataset_scores(model_report)
                values = [(scores.get(c) or 0) / 100.0 for c in dataset_codes]
                radii = [max_r * v for v in values]
                if num_metrics == 1:
                    xs = [0, radii[0] * math.cos(vertex_angles[0])]
                    ys = [0, radii[0] * math.sin(vertex_angles[0])]
                    ax.plot(xs, ys, 'o-', linewidth=2, label=model_report.model_name, color=colors[idx % len(colors)])
                else:
                    poly_angles = vertex_angles + [vertex_angles[0]]
                    poly_radii = radii + [radii[0]]
                    xs = [poly_radii[i] * math.cos(poly_angles[i]) for i in range(len(poly_angles))]
                    ys = [poly_radii[i] * math.sin(poly_angles[i]) for i in range(len(poly_angles))]
                    ax.plot(xs, ys, 'o-', linewidth=2, label=model_report.model_name, color=colors[idx % len(colors)])
                    ax.fill(xs, ys, alpha=0.15, color=colors[idx % len(colors)])
            label_r = max_r * 1.08
            for i in range(num_metrics):
                a = vertex_angles[i]
                ax.text(label_r * math.cos(a), label_r * math.sin(a), metric_labels[i], ha='center', va='center', fontsize=10)
            ax.set_xlim(-1.25 * max_r, 1.25 * max_r)
            ax.set_ylim(-1.25 * max_r, 1.25 * max_r)
            ax.set_aspect('equal')
            ax.axis('off')
            ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.0))
            plt.tight_layout()
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成对比评估雷达图失败: {str(e)}", exc_info=True)
            plt.close()
            return None

    def _add_bar_charts(self):
        """第三部分：评分对比柱状图。x 轴为评估指标，y 轴为得分，不同颜色柱体代表不同模型。"""
        self.doc.add_heading('三、评分对比柱状图', 1)
        if not self.compare_data.model_reports:
            self.doc.add_paragraph("暂无评估结果数据。")
            return
        dataset_codes, metric_labels = self._get_dataset_codes_and_labels()
        if not dataset_codes:
            self.doc.add_paragraph("暂无评估指标数据。")
            return
        try:
            bar_image = self._generate_comparison_bar_chart(dataset_codes, metric_labels)
            if bar_image:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(BytesIO(bar_image), width=Inches(6))
        except Exception as e:
            logger.warning(f"生成对比评估柱状图失败: {str(e)}")
            self.doc.add_paragraph(f"[柱状图生成失败：{str(e)}]")

    def _generate_comparison_bar_chart(self, dataset_codes: List[str], metric_labels: List[str]) -> Optional[bytes]:
        """柱状图：横轴=评估指标，纵轴=得分，分组为各模型。"""
        if not dataset_codes or not self.compare_data.model_reports:
            return None
        try:
            model_names = [r.model_name for r in self.compare_data.model_reports]
            x = range(len(dataset_codes))
            width = 0.8 / len(model_names)
            colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', '#13c2c2', '#eb2f96']
            fig, ax = plt.subplots(figsize=(12, 6))
            max_val = 0
            for i, model_report in enumerate(self.compare_data.model_reports):
                scores_map = self._model_report_to_dataset_scores(model_report)
                scores = [scores_map.get(c) or 0 for c in dataset_codes]
                max_val = max(max_val, max(scores) if scores else 0)
                offset = (i - len(model_names) / 2 + 0.5) * width
                bars = ax.bar(
                    [xi + offset for xi in x],
                    scores,
                    width,
                    label=model_report.model_name,
                    color=colors[i % len(colors)],
                )
                for bar in bars:
                    h = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2.0, h, f'{h:.2f}', ha='center', va='bottom', fontsize=8)
            ax.set_ylabel('得分', fontsize=12)
            ax.set_xticks(x)
            ax.set_xticklabels(metric_labels, rotation=45, ha='right')
            ax.set_ylim(0, max(max_val * 1.2, 100) if max_val > 0 else 100)
            ax.legend()
            plt.tight_layout()
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"生成对比评估柱状图失败: {str(e)}", exc_info=True)
            plt.close()
            return None
