"""
DOCX 导出报告正文字体：统一东亚字体（Word 中显示为「黑体」）。
图表由 matplotlib 单独配置字体，不使用本模块。
"""
from __future__ import annotations

from typing import TYPE_CHECKING, Iterator

from docx.oxml.ns import qn

from app.core.logging import logger

if TYPE_CHECKING:
    from docx import Document
    from docx.table import _Cell
    from docx.text.run import Run

# Word 中文字体名（Windows 常见为「黑体」，对应西文槽可同名或 SimHei）
DOCX_BODY_FONT_CN = "黑体"


def setup_document_east_asia_font(document: "Document", font_name: str = DOCX_BODY_FONT_CN) -> None:
    """将文档各样式的默认字体设为指定中文字体（含东亚字体槽）。"""
    for style in document.styles:
        if not hasattr(style, "font"):
            continue
        try:
            style.font.name = font_name
            el = style._element
            if el.rPr is None or el.rPr.rFonts is None:
                continue
            el.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        except (AttributeError, TypeError) as e:
            # 如果出现样式设置异常，仍执行后续代码，避免因样式设置异常导致整个导出过程失败
            style_id = getattr(style, "name", None) or getattr(style, "style_id", None) or type(style).__name__
            logger.warning(
                "设置文档样式东亚字体跳过: style=%r font=%r err=%s",
                style_id,
                font_name,
                e,
            )
            continue


def set_run_east_asia_font(run: "Run", font_name: str = DOCX_BODY_FONT_CN) -> None:
    """单个 Run：西文/东亚字体槽均指向黑体，避免标题等单独改字号后仍用宋体/Calibri。"""
    run.font.name = font_name
    try:
        r_pr = run._element.rPr
        if r_pr is None:
            return
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            return
        r_fonts.set(qn("w:eastAsia"), font_name)
    except (AttributeError, TypeError) as e:
        # 如果设置字体出现异常，仍执行后续代码，避免因字体异常导致整个导出过程失败
        logger.debug(
            "设置 Run 东亚字体跳过: font=%r err=%s",
            font_name,
            e,
        )


def _runs_in_cell(cell: "_Cell") -> Iterator["Run"]:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            yield run
    for table in getattr(cell, "tables", ()) or ():
        for row in table.rows:
            for c in row.cells:
                yield from _runs_in_cell(c)


def apply_east_asia_font_to_all_runs(document: "Document", font_name: str = DOCX_BODY_FONT_CN) -> None:
    """
    遍历正文段落与表格（含嵌套表）中所有 Run，强制东亚字体。
    插入的图片/绘图 Run 设置字体不影响图本身，可一并处理。
    """
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            set_run_east_asia_font(run, font_name)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for run in _runs_in_cell(cell):
                    set_run_east_asia_font(run, font_name)
